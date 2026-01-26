"""
DocuMind AI - Word 文档解析器

使用 python-docx 解析 DOCX 文件
"""

from pathlib import Path
from typing import Any, Dict

from .base import BaseParser, ParseResult


class DocxParser(BaseParser):
    """Word 文档解析器"""

    supported_extensions = ["docx"]
    name = "DocxParser"

    def parse(self, file_path: str) -> ParseResult:
        """
        解析 Word 文档

        Args:
            file_path: DOCX 文件路径

        Returns:
            ParseResult 解析结果
        """
        # 验证文件
        error = self._validate_file(file_path)
        if error:
            return self._create_error_result(error)

        try:
            from docx import Document

            # 打开文档
            doc = Document(file_path)

            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            content = "\n\n".join(paragraphs)
            content = self._clean_text(content)

            # 提取元数据
            metadata = self._extract_metadata(doc, file_path)

            return ParseResult(
                content=content,
                metadata=metadata,
                success=True,
            )

        except ImportError:
            return self._create_error_result(
                "缺少 python-docx 依赖，请运行: pip install python-docx"
            )
        except Exception as e:
            return self._create_error_result(f"DOCX 解析错误: {str(e)}")

    def _extract_metadata(self, doc, file_path: str) -> Dict[str, Any]:
        """提取 Word 文档元数据"""
        path = Path(file_path)

        metadata = {
            "filename": path.name,
            "file_size": path.stat().st_size,
            "format": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        # 文档核心属性
        core_props = doc.core_properties
        if core_props:
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["creation_date"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified_date"] = core_props.modified.isoformat()

        return metadata
